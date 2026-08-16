# -*- coding: utf-8 -*-
"""从 data/resume.json 生成正式 DOCX 简历（A4、中文排版、可编辑、URL 可点击）。
用法: py scripts/gen_docx.py
输出: public/秦宇-测试工程师.docx 与 docs/秦宇-测试工程师.docx
"""
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(ROOT, "data", "resume.json"), encoding="utf-8") as f:
    R = json.load(f)

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x05, 0x63, 0xC1)  # 链接色（Word 默认超链接蓝）


def set_font(run, name_cn="微软雅黑", name_en="Calibri", size=10.5, bold=False, color=None):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)


def add_hyperlink(paragraph, url, text, size=10.5):
    """可点击超链接（真实 URL，ATS 可见）。"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = run.makeelement(qn("w:rPr"), {})
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.append(rFonts)
    color = rPr.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    rPr.append(color)
    u = rPr.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rPr.append(u)
    run.append(rPr)
    t = run.makeelement(qn("w:t"), {})
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def keep_together(p):
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False


def main():
    doc = Document()
    # 页面：A4，2cm 页边距
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Cm(1.9))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    b = R["basics"]

    # 姓名 + 头衔
    p = doc.add_paragraph()
    set_font(p.add_run(b["name"]), size=26, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    set_font(p.add_run(f"{b['brand']['primary']} · {' · '.join(b['brand']['subs'])}"), size=11, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(4)

    # 联系方式（真实文本，ATS 可解析）
    p = doc.add_paragraph()
    set_font(p.add_run("电话 "), size=10, color=MUTED)
    set_font(p.add_run(b["contact"]["phone"]), size=10)
    set_font(p.add_run(" ｜ 邮箱 "), size=10, color=MUTED)
    add_hyperlink(p, f"mailto:{b['contact']['email']}", b["contact"]["email"])
    set_font(p.add_run(" ｜ 本科 · 软件工程 · 1 年经验 ｜ 期望城市："), size=10, color=MUTED)
    set_font(p.add_run(b["contact"]["location"]), size=10)
    set_font(p.add_run(" ｜ 期望薪资："), size=10, color=MUTED)
    set_font(p.add_run(b["contact"]["salary"]), size=10)
    set_font(p.add_run(" ｜ "), size=10, color=MUTED)
    set_font(p.add_run(b["contact"]["availability"]), size=10)
    p.paragraph_format.space_after = Pt(3)
    p = doc.add_paragraph()
    set_font(p.add_run("GitHub："), size=10, color=MUTED)
    add_hyperlink(p, b["contact"]["github"], b["contact"]["github"])
    p.paragraph_format.space_after = Pt(8)
    # 分割线
    pPr = doc.add_paragraph()._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "12", qn("w:color"): "1A1A1A"})
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 定位
    p = doc.add_paragraph()
    set_font(p.add_run("职业定位："), bold=True, size=10.5)
    set_font(p.add_run(
        f"{b['positioning']} 现于半导体芯片测试机上市公司负责 ATE 上位机软件测试（黑盒 + UI 自动化）；"
        f"此前在 AI 智能体与低代码平台独立完成 1000+ 用例与 50+ 核心 API 自动巡检。"
    ))
    keep_together(p)

    def section(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(title), size=12.5, bold=True, color=INK)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        left = pPr.makeelement(qn("w:left"), {qn("w:val"): "single", qn("w:sz"): "18", qn("w:color"): "1A1A1A"})
        pBdr.append(left)
        pPr.append(pBdr)

    def bullet(lead, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run("▪ "), size=10)
        if lead:
            set_font(p.add_run(f"{lead}："), bold=True)
        set_font(p.add_run(text))
        keep_together(p)

    # 核心亮点
    section("核心亮点")
    for h in R["highlights"]:
        bullet("", f"{h['value']} {h['label']}（{h['sub']}）")

    # 专业技能（能力地图）
    section("专业技能")
    for cap in R["capabilities"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"{cap['name']} "), bold=True)
        set_font(p.add_run(f"（{cap['level']}） "), size=9.5, color=MUTED)
        set_font(p.add_run(cap["techs"].join(" / ") if isinstance(cap["techs"], str) else " / ".join(cap["techs"])))
        keep_together(p)

    # 工作经历
    section("工作经历")
    for exp in R["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run(exp["company"] + "　"), size=11.5, bold=True)
        set_font(p.add_run(f"[{exp['role']}]　"), size=9.5, color=MUTED)
        set_font(p.add_run(exp["period"]), size=9.5, color=MUTED)
        keep_together(p)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(exp["industry"]), size=9.5, color=MUTED)
        for b2 in exp["bullets"]:
            bullet(b2["lead"], b2["text"])

    # 项目经历
    section("项目经历")
    for pr in R["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run(f"{pr['title']}　"), size=11, bold=True)
        set_font(p.add_run(f"[{pr['badge']}]　"), size=9.5, color=MUTED)
        set_font(p.add_run(pr["period"]), size=9.5, color=MUTED)
        keep_together(p)
        bullet("问题与方法", f"{pr['problem']} {pr['approach']}")
        for s in pr["strategy"]:
            bullet("", s)
        bullet("结果", pr["result"])

    # 教育
    section("教育 · 荣誉 · 证书")
    edu = R["education"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{edu['school']}｜{edu['major']}｜{edu['period']}"), bold=True)
    p = doc.add_paragraph()
    set_font(p.add_run(edu["extras"]), size=9.5, color=MUTED)

    # 输出
    os.makedirs(os.path.join(ROOT, "public"), exist_ok=True)
    os.makedirs(os.path.join(ROOT, "docs"), exist_ok=True)
    out_main = os.path.join(ROOT, "public", "秦宇-测试工程师.docx")
    out_docs = os.path.join(ROOT, "docs", "秦宇-测试工程师.docx")
    doc.save(out_main)
    doc.save(out_docs)
    print(f"DOCX 已生成: {out_main}")
    print(f"DOCX 已生成: {out_docs}")


if __name__ == "__main__":
    sys.exit(main())
